"""
Export all email templates for the AI Make More Money & Freedom event
into a Word document for review/edit by Lê Đăng Khương.

This version (v2):
- DROPS specific calendar dates from email bodies — uses "Buổi 1/2/3"
  and relative timing ("tối mai", "còn 1 tiếng") so the same templates
  can be reused for every cohort.
- ADDS the program-overview block (curriculum bullets for all 3 buổi)
  to welcome + D-1 emails so customers know what they'll learn.
- ADDS an upgrade CTA block for FREE (→ VIP + VVIP) and VIP (→ VVIP).
- CLARIFIES that {tên_khách} in this doc is just a marker — the real
  email injects the customer name via the `name` parameter and contains
  no `{name}` placeholder text.
"""
import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
from docx.shared import RGBColor, Pt, Inches

OUTPUT = r"C:\Users\Admin\Downloads\AI-Make-More-Money-Email-Automation-v2.docx"

# ── Color palette ──
GOLD = RGBColor(0xD4, 0xA8, 0x43)
GREEN = RGBColor(0x22, 0xC5, 0x5E)
GRAY = RGBColor(0x6B, 0x72, 0x80)
RED = RGBColor(0xEF, 0x44, 0x44)
BLUE = RGBColor(0x00, 0x68, 0xFF)
BLACK = RGBColor(0x0A, 0x0A, 0x0A)

doc = Document()
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

def H1(text, color=GOLD):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = color
    return p

def H2(text, color=GOLD):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(15)
    run.font.color.rgb = color
    return p

def H3(text, color=BLACK):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = color
    return p

def label(text, color=GRAY):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(9)
    run.font.color.rgb = color
    p.paragraph_format.space_after = Pt(2)
    return p

def body(text):
    p = doc.add_paragraph()
    p.add_run(text)
    return p

def callout(title, content, color=GOLD):
    p = doc.add_paragraph()
    run = p.add_run(f"📌 {title}\n")
    run.bold = True
    run.font.color.rgb = color
    p.add_run(content)
    p.paragraph_format.left_indent = Inches(0.2)
    return p

def divider():
    doc.add_paragraph("─" * 60)

def bullet(text):
    return doc.add_paragraph(text, style='List Bullet')

# ── Reusable curriculum block ──
SESSIONS = {
    1: {
        "title": "Tư Duy Đúng & 10 Nguồn Thu Nhập Từ AI 2026",
        "bullets": [
            "Tư duy đúng để kiếm tiền bằng AI — vì sao đây là thời điểm vàng",
            "Toàn cảnh 10 nguồn thu nhập đến từ AI và bạn nên bắt đầu từ đâu",
            "Lộ trình kiếm 10 nguồn thu nhập trên internet, từ con số 0",
        ],
    },
    2: {
        "title": "Video & Kênh Triệu View — Kiếm Tiền Từ Affiliate",
        "bullets": [
            "Cách tạo video AI hấp dẫn và xây kênh triệu view, không cần quay dựng",
            "Kiếm tiền Affiliate ở 4 ngách hot: KOL AI · Tiếng Anh · Sức khỏe · Sách",
            "Công thức biến lượt xem thành hoa hồng đều đặn",
        ],
    },
    3: {
        "title": "Chuyển Đổi Khách Thành Tiền & Hệ Thống Tự Động",
        "bullets": [
            "Bí mật chuyển đổi danh sách khách hàng thành tiền",
            "Cách xây dựng sản phẩm số từ chính chuyên môn của bạn",
            "Dựng Website All-in-One bán hàng tự động với AI Agent — bạn ngủ, hệ thống vẫn bán",
        ],
    },
}

def session_block(n, label_text=None):
    label(f"BUỔI {n} · 20:00 – 22:00")
    body(SESSIONS[n]["title"])
    for b in SESSIONS[n]["bullets"]:
        bullet(b)
    body("")

def program_overview_block():
    label("🎯 NỘI DUNG 3 BUỔI (xuất hiện trong welcome + D-1 reminder)")
    for n in [1, 2, 3]:
        session_block(n)

def upgrade_block_free():
    label("🚀 BLOCK NÂNG CẤP (cho vé FREE — hiện ở welcome, D-1, recap)")
    body("Tiêu đề: Nâng cấp để giữ trọn giá trị")
    body("Mô tả: Vé Free chỉ học trực tiếp, không có video xem lại. Nếu bỏ lỡ buổi nào, bạn sẽ mất buổi đó.")
    body("")
    body("👉 Vé VIP — 99.000đ (≈ 1 ly cà phê):")
    bullet("Video xem lại vĩnh viễn cả 3 buổi")
    bullet("Bộ slide PDF + tài liệu từng buổi")
    bullet("Ưu tiên đặt câu hỏi Q&A")
    body("→ Nút (vàng): 'Nâng cấp VIP — 99k'  ·  Link: /aimakemoremoney#tickets")
    body("")
    body("👉 Vé VVIP — 499.000đ (giới hạn 50 suất):")
    bullet("Toàn bộ quyền lợi VIP")
    bullet("30 phút coaching 1-1 trực tiếp với Khương")
    bullet("AI Agent Starter Kit (template + prompt + hướng dẫn)")
    body("→ Nút (xanh): 'Xem chi tiết VVIP — 499k'  ·  Link: /aimakemoremoney#tickets")
    body("")

def upgrade_block_vip():
    label("💎 BLOCK NÂNG CẤP VVIP (cho vé VIP — hiện ở welcome, D-1, recap)")
    body("Tiêu đề: Muốn được Khương kèm 1-1?")
    body("Mô tả: Học viên VIP có quyền nâng cấp lên VVIP — bao gồm coaching 1-1 + AI Agent Starter Kit.")
    body("")
    bullet("30 phút coaching 1-1 trực tiếp với Lê Đăng Khương — phân tích case riêng của bạn")
    bullet("AI Agent Starter Kit — template database + prompt library + hướng dẫn dựng AI Agent đầu tiên")
    bullet("Ưu tiên Q&A số 1 — câu hỏi của bạn được trả lời trước")
    body("→ Nút (xanh): 'Nâng cấp VVIP — 499k'  ·  Link: /aimakemoremoney#tickets")
    body("")

# ════════════════════════════════════════════════════════════════
# COVER PAGE
# ════════════════════════════════════════════════════════════════
H1("🚀 AI MAKE MORE MONEY & FREEDOM", color=GOLD)
H2("Hệ thống Email Automation — Bản nội dung để review (v2)", color=BLACK)
body("")
label("LANDING PAGE")
body("https://dangkhuong.com/aimakemoremoney")
label("ZALO GROUP")
body("https://zalo.me/g/l4qmpdq934rmst9xxnfj")
label("LỊCH SỰ KIỆN (config trong cron)")
body("3 buổi liên tiếp · 20:00 – 22:00 mỗi buổi · Zoom")
body("Anh chạy event nhiều cohort — nội dung email KHÔNG nhắc ngày cụ thể, "
     "chỉ nói 'Buổi 1/2/3', 'tối mai', 'còn 1 tiếng'. "
     "Ngày thực tế chỉ có trong file cron (schedule) — anh update khi mở cohort mới.")
body("")

callout(
    "Quan trọng — placeholder tên khách",
    "Trong tài liệu này em viết {tên_khách} để đánh dấu nơi tên thật được chèn vào. "
    "Email thật KHÔNG có chữ '{tên_khách}' — code tự inject tên thật từ database. "
    "Ví dụ khách tên 'Trần Thế Anh' sẽ nhận email mở đầu: 'Chào Trần Thế Anh,...'"
)
body("")

callout(
    "Cách dùng tài liệu",
    "Mỗi email có 3 biến thể (FREE / VIP / VVIP). Anh review từng email, sửa text "
    "trực tiếp lên file Word. Sau đó gửi lại em qua chat — em sẽ cập nhật code và "
    "push deploy. Phần code (link, nút, cấu trúc) anh không cần sửa.",
    color=BLUE,
)

body("")
label("📅 TỔNG QUAN — LỊCH TRÌNH GỬI 11 EMAIL")
schedule_table = doc.add_table(rows=12, cols=3)
schedule_table.style = 'Light Grid Accent 4'
schedule_table.rows[0].cells[0].text = "Email"
schedule_table.rows[0].cells[1].text = "Thời điểm gửi"
schedule_table.rows[0].cells[2].text = "Mô tả ngắn"

schedule = [
    ("1. Welcome", "Ngay khi đăng ký", "Xác nhận + tổng quan 3 buổi + Zalo + upgrade CTA"),
    ("2. D-1 Buổi 1", "Trưa 1 ngày trước Buổi 1", "Nhắc tối mai 20:00 + nội dung Buổi 1"),
    ("3. D-1 Buổi 2", "Trưa 1 ngày trước Buổi 2", "Nhắc tối mai 20:00 + nội dung Buổi 2"),
    ("4. D-1 Buổi 3", "Trưa 1 ngày trước Buổi 3", "Nhắc tối mai 20:00 + nội dung Buổi 3"),
    ("5. T-1h Buổi 1", "19:00 ngày diễn ra Buổi 1", "Còn 1h nữa + checklist 5 phút"),
    ("6. T-1h Buổi 2", "19:00 ngày diễn ra Buổi 2", "Còn 1h nữa + checklist"),
    ("7. T-1h Buổi 3", "19:00 ngày diễn ra Buổi 3", "Còn 1h nữa + checklist"),
    ("8. Recap Buổi 1", "22:30 ngày diễn ra Buổi 1", "Cảm ơn + replay/upsell + nhắc Buổi 2"),
    ("9. Recap Buổi 2", "22:30 ngày diễn ra Buổi 2", "Cảm ơn + nhắc Buổi 3"),
    ("10. Recap Buổi 3", "22:30 ngày diễn ra Buổi 3", "Cảm ơn + dẫn về event complete"),
    ("11. Event Complete", "Sáng D+1 (ngày sau buổi cuối)", "Bước tiếp theo + upsell theo tier"),
]
for i, (n, t, note) in enumerate(schedule):
    r = schedule_table.rows[i + 1]
    r.cells[0].text = n
    r.cells[1].text = t
    r.cells[2].text = note

doc.add_page_break()

# ════════════════════════════════════════════════════════════════
# EMAIL 1: WELCOME (3 tier variants)
# ════════════════════════════════════════════════════════════════
H1("📧 EMAIL 1: WELCOME", color=GOLD)
label("THỜI ĐIỂM: Ngay khi khách bấm đăng ký")
body("")

# ── Tier: FREE ──
H2("Biến thể 1.1: Vé FREE", color=GRAY)
label("SUBJECT")
body("🎉 Chào mừng {tên_khách} — đăng ký AI Make More Money & Freedom thành công")
label("BODY")
body("Chào {tên_khách}, chào mừng đến với AI Make More Money & Freedom 🚀")
body("")
body("Vé Free cho bạn quyền tham gia trực tiếp cả 3 buổi Zoom + nhận quà cẩm nang trị giá 2.990.000đ.")
body("")
body("⚠️ Lưu ý: vé Free không có video xem lại — nếu bỏ lỡ buổi nào, bạn sẽ mất buổi đó.")
body("")
label("🎯 TỔNG QUAN CHƯƠNG TRÌNH (hiện ngay sau lời mở thư)")
body("3 buổi tối Zoom, mỗi buổi 20:00 – 22:00. Bạn sẽ đi qua trọn vẹn lộ trình: "
     "Tư duy đúng → Sản xuất nội dung → Chuyển đổi & tự động hoá — kiếm tiền bằng AI "
     "mà không cần giỏi công nghệ, không cần làm nhiều hơn.")
body("")
program_overview_block()
label("💬 BLOCK ZALO CALLOUT")
body("Tiêu đề: Tham gia nhóm Zalo (quan trọng!)")
body("Nội dung: Link Zoom 3 buổi sẽ được gửi qua nhóm Zalo trước mỗi buổi 30 phút.")
body("→ Nút (xanh Zalo): 'Vào nhóm Zalo ngay'  ·  Link: zalo.me/g/l4qmpdq934rmst9xxnfj")
body("")
upgrade_block_free()
label("CTA CUỐI THƯ")
body("→ Nút (vàng): 'Vào trang khoá học của tôi'  ·  Link: /courses/ai-make-more-money-free")
body("")
label("KẾT THƯ")
body("Có vấn đề gì? Trả lời thẳng email này — em sẽ hỗ trợ trong 24h.")

divider()

# ── Tier: VIP ──
H2("Biến thể 1.2: Vé VIP", color=GOLD)
label("SUBJECT (giống Free)")
body("🎉 Chào mừng {tên_khách} — đăng ký AI Make More Money & Freedom thành công")
label("BODY (giới thiệu)")
body("Chào {tên_khách}, chào mừng đến với AI Make More Money & Freedom 🚀")
body("")
body("Vé VIP của bạn đã được kích hoạt. Bạn có quyền truy cập video xem lại VĨNH VIỄN + "
     "bộ slide PDF từng buổi + ưu tiên đặt câu hỏi trong Q&A.")
body("")
body("Sau mỗi buổi học, video replay sẽ được cập nhật trong vòng 24h vào trang khoá học của bạn.")
body("")
label("→ Sau đó hiện 'TỔNG QUAN CHƯƠNG TRÌNH + 3 BUỔI' giống Free")
body("→ Sau đó hiện ZALO CALLOUT giống Free")
body("")
upgrade_block_vip()
label("CTA CUỐI THƯ")
body("→ Nút (vàng): 'Vào trang khoá học của tôi'  ·  Link: /courses/ai-make-more-money-vip")

divider()

# ── Tier: VVIP ──
H2("Biến thể 1.3: Vé VVIP", color=GREEN)
label("SUBJECT (giống Free/VIP)")
body("🎉 Chào mừng {tên_khách} — đăng ký AI Make More Money & Freedom thành công")
label("BODY")
body("Chào {tên_khách}, chào mừng đến với AI Make More Money & Freedom 🚀")
body("")
body("Vé VVIP — suất gần Thầy Khương nhất — đã được kích hoạt. Bạn có toàn bộ quyền lợi của VIP "
     "+ 30 phút coaching 1-1 trực tiếp với Lê Đăng Khương + AI Agent Starter Kit.")
body("")
body("📅 Link đặt lịch coaching 1-1 sẽ được gửi qua email riêng sau Buổi 3. "
     "Hãy giữ lịch trống tuần sau!")
body("")
label("→ Sau đó hiện 'TỔNG QUAN CHƯƠNG TRÌNH + 3 BUỔI' giống Free")
body("→ Sau đó hiện ZALO CALLOUT giống Free")
body("→ VVIP KHÔNG có upgrade block (đã tier cao nhất)")
body("")
label("CTA CUỐI THƯ")
body("→ Nút (vàng): 'Vào trang khoá học của tôi'  ·  Link: /courses/ai-make-more-money-vvip")

doc.add_page_break()

# ════════════════════════════════════════════════════════════════
# EMAIL 2-4: D-1 REMINDER
# ════════════════════════════════════════════════════════════════
H1("📧 EMAIL 2-4: D-1 REMINDER (nhắc 1 ngày trước mỗi buổi)", color=GOLD)
label("THỜI ĐIỂM")
body("• Email 2: trưa 1 ngày trước Buổi 1")
body("• Email 3: trưa 1 ngày trước Buổi 2")
body("• Email 4: trưa 1 ngày trước Buổi 3")
body("")
callout(
    "Lưu ý",
    "Cả 3 email có cùng cấu trúc — chỉ khác số buổi (1/2/3) và tên buổi. "
    "Tier (FREE/VIP/VVIP) khác ở badge đầu thư + có/không có upgrade block. "
    "Em show mẫu Email 2 (Buổi 1), Email 3/4 chỉ cần thay số/tên."
)

H2("Mẫu — D-1 Buổi 1", color=GRAY)
label("SUBJECT")
body("⏰ Tối mai 20:00 — Buổi 1: Tư Duy Đúng & 10 Nguồn Thu Nhập Từ AI 2026")
label("BODY")
body("{tên_khách}, tối mai 20:00 mình gặp nhau ở Zoom 🎯")
body("")
body("Buổi 1 / 3 của chương trình AI Make More Money & Freedom diễn ra tối mai, 20:00 – 22:00.")
body("")
label("📌 CARD BUỔI 1 (highlight vàng — nổi bật ngay sau lời mở)")
session_block(1)
label("📅 NỘI DUNG 3 BUỔI (chỉ hiện ở D-1 Buổi 1 & 2 — Buổi 3 không cần)")
body("Để bạn nắm trọn lộ trình, đây là toàn bộ 3 buổi:")
program_overview_block()
label("📝 CHUẨN BỊ TRƯỚC BUỔI HỌC")
bullet("Notebook + bút để ghi chú")
bullet("Máy tính có kết nối Internet ổn định")
bullet("Tinh thần học hỏi mở & sẵn sàng áp dụng ngay")
body("")
label("💬 BLOCK ZALO CALLOUT (chung)")
body("→ Nút (xanh Zalo): 'Vào nhóm Zalo ngay'")
body("")
label("🚀 UPGRADE BLOCK (chỉ cho FREE & VIP, VVIP không có)")
body("→ Free thấy block VIP+VVIP upgrade")
body("→ VIP thấy block VVIP upgrade")
body("→ VVIP không hiện block")
body("")
label("KẾT THƯ")
body("Link Zoom sẽ được gửi qua nhóm Zalo + email TRƯỚC buổi 30 phút. Đảm bảo bạn đã vào nhóm Zalo!")

divider()

H3("Email 3 — D-1 Buổi 2")
label("SUBJECT")
body("⏰ Tối mai 20:00 — Buổi 2: Video & Kênh Triệu View — Kiếm Tiền Từ Affiliate")
body("→ Cấu trúc giống Email 2 nhưng đổi số → 'Buổi 2 / 3' và highlight Card Buổi 2")

H3("Email 4 — D-1 Buổi 3 (buổi cuối)")
label("SUBJECT")
body("⏰ Tối mai 20:00 — Buổi 3: Chuyển Đổi Khách Thành Tiền & Hệ Thống Tự Động")
body("→ Khác Email 2/3: KHÔNG hiện 'Nội dung 3 buổi' nữa vì đã đến buổi cuối")
body("→ Chỉ hiện Card Buổi 3 + chuẩn bị + Zalo + upgrade (nếu free/vip)")

doc.add_page_break()

# ════════════════════════════════════════════════════════════════
# EMAIL 5-7: T-1h REMINDER
# ════════════════════════════════════════════════════════════════
H1("📧 EMAIL 5-7: T-1h REMINDER (nhắc 1 giờ trước mỗi buổi)", color=GOLD)
label("THỜI ĐIỂM")
body("• Email 5: 19:00 ngày diễn ra Buổi 1")
body("• Email 6: 19:00 ngày diễn ra Buổi 2")
body("• Email 7: 19:00 ngày diễn ra Buổi 3")
body("")

H2("Mẫu — T-1h Buổi 1", color=GRAY)
label("SUBJECT")
body("🔥 Còn 1 tiếng nữa — Buổi 1: Tư Duy Đúng & 10 Nguồn Thu Nhập Từ AI 2026")
label("BODY")
body("{tên_khách}, còn 1 tiếng nữa — Buổi 1 bắt đầu! ⚡")
body("")
body("Đúng 20:00 tối nay, mình sẽ cùng nhau khai mạc Buổi 1.")
body("")
label("📌 CARD BUỔI 1 (highlight vàng — gồm cả 3 bullets nội dung)")
session_block(1)
label("✅ CHECKLIST 5 PHÚT TRƯỚC BUỔI (box xanh)")
bullet("Vào nhóm Zalo — link Zoom được pin ở đầu nhóm")
bullet("Tắt thông báo & chuẩn bị nơi yên tĩnh")
bullet("Test mic + camera Zoom")
bullet("Có sẵn notebook để ghi chú")
body("")
label("CTA chính")
body("→ Nút (xanh Zalo): 'Vào nhóm Zalo lấy link Zoom'")
body("→ Link: zalo.me/g/l4qmpdq934rmst9xxnfj")
body("")
label("Lưu ý")
body("Email T-1h KHÔNG có upgrade block — focus 100% vào việc giúp khách vào học đúng giờ.")

divider()
body("Email 6 (Buổi 2) và Email 7 (Buổi 3) cùng cấu trúc — chỉ thay số buổi & tên.")

doc.add_page_break()

# ════════════════════════════════════════════════════════════════
# EMAIL 8-10: RECAP
# ════════════════════════════════════════════════════════════════
H1("📧 EMAIL 8-10: RECAP (sau mỗi buổi, gửi 22:30)", color=GOLD)
label("THỜI ĐIỂM")
body("• Email 8: 22:30 ngày diễn ra Buổi 1")
body("• Email 9: 22:30 ngày diễn ra Buổi 2")
body("• Email 10: 22:30 ngày diễn ra Buổi 3 (buổi cuối — kết thư khác)")
body("")
callout(
    "Khác biệt theo tier",
    "Recap email khác nhau ở phần REPLAY BLOCK. "
    "FREE → upsell VIP. VIP/VVIP → link đến replay khoá học. "
    "Cảm ơn + nhắc buổi sau (nếu có) thì chung."
)

H2("Mẫu — Recap Buổi 1", color=GRAY)
label("SUBJECT (chung cả 3 tier)")
body("💎 Cảm ơn {tên_khách} — Buổi 1 đã kết thúc, hẹn buổi tiếp theo")
label("BODY (chung)")
body("Cảm ơn {tên_khách} đã tham gia Buổi 1 🙏")
body("")
body('Hi vọng bạn đã có những insight giá trị từ buổi "Tư Duy Đúng & 10 Nguồn Thu Nhập Từ AI 2026".')
body("")

H3("REPLAY BLOCK — Vé FREE (đỏ, upsell)", color=RED)
body("📺 Bỏ lỡ buổi? Muốn xem lại?")
body("Vé Free không có replay. Nâng cấp lên VIP (99k) để có video xem lại vĩnh viễn cả 3 buổi + slide PDF.")
body("→ Nút (vàng): 'Nâng cấp VIP — chỉ 99k'  ·  Link: /aimakemoremoney#tickets")
body("")

H3("REPLAY BLOCK — Vé VIP & VVIP (xanh, link replay)", color=GREEN)
body("🎬 Video xem lại Buổi 1")
body("Replay sẽ được cập nhật vào trang khoá học của bạn TRONG VÒNG 24H. "
     "Bạn xem lại vĩnh viễn — học bao nhiêu lần tuỳ thích.")
body("→ Nút (vàng): 'Mở khoá học của tôi'  ·  Link: /courses/ai-make-more-money-{tier}")
body("")

label("👉 NHẮC BUỔI TIẾP THEO (chỉ Recap Buổi 1 & 2, KHÔNG có ở Buổi 3)")
session_block(2)
label("UPGRADE BLOCK (chỉ Free & VIP, KHÔNG ở Recap Buổi 3 vì sẽ có ở Event Complete)")
body("→ Free thấy upgrade VIP/VVIP")
body("→ VIP thấy upgrade VVIP")
body("→ VVIP không hiện")

divider()

H3("Email 10 — Recap Buổi 3 (buổi cuối)")
label("SUBJECT (khác Buổi 1, 2)")
body("🎉 Đã xong 3 buổi — cảm ơn {tên_khách} đã đồng hành!")
label("KẾT THƯ thay vì 'Buổi tiếp theo'")
body("🚀 Bắt đầu hành trình của bạn ngay tuần này!")
body("3 buổi đã cho bạn bản đồ. Bước tiếp theo là ÁP DỤNG. Em hẹn gặp bạn trong cộng đồng "
     "— chia sẻ kết quả nhé!")
body("→ Nút: 'Vào dashboard của tôi'  ·  Link: /dashboard")

doc.add_page_break()

# ════════════════════════════════════════════════════════════════
# EMAIL 11: EVENT COMPLETE
# ════════════════════════════════════════════════════════════════
H1("📧 EMAIL 11: EVENT COMPLETE (sáng D+1)", color=GOLD)
label("MỤC ĐÍCH")
body("• Tổng kết sự kiện và push khách về bước tiếp theo")
body("• 3 tier có nội dung HOÀN TOÀN khác nhau (upsell theo từng nhóm)")
body("")
label("SUBJECT (chung cả 3 tier)")
body("🎁 {tên_khách}, đây là bước tiếp theo của bạn sau AI Make More Money")
body("")
label("MỞ THƯ (chung)")
body("Chào {tên_khách} — chương trình đã khép lại, nhưng hành trình của bạn VỪA BẮT ĐẦU ✨")
body("")
body("3 buổi vừa qua đã trao cho bạn bản đồ kiếm tiền bằng AI. "
     "Bước quan trọng nhất giờ là ÁP DỤNG NGAY TUẦN NÀY — đừng để kiến thức nằm im trên giấy.")
body("")

divider()

H2("Biến thể 11.1: Vé FREE — Upsell VIP + VVIP", color=RED)
label("PHẦN HEADING")
body("🚀 Muốn xem lại 3 buổi vừa qua?")
body("Vé Free đã hết quyền truy cập video. Nếu bạn muốn xem lại + có slide PDF "
     "để học sâu hơn, hãy nâng cấp lên VIP.")
body("")
label("UPGRADE BLOCK (đầy đủ cả VIP và VVIP)")
upgrade_block_free()

divider()

H2("Biến thể 11.2: Vé VIP — Replay + Upsell VVIP", color=GOLD)
label("PHẦN CONFIRM REPLAY")
body("🎬 Toàn bộ replay đã sẵn sàng")
body("Cả 3 buổi đã được upload đầy đủ vào trang khoá học VIP của bạn. "
     "Xem lại bất cứ lúc nào, slide PDF cũng có sẵn để tải.")
body("→ Nút (vàng, ở giữa): 'Mở khoá VIP của tôi'  ·  Link: /courses/ai-make-more-money-vip")
body("")
label("UPGRADE BLOCK VVIP")
upgrade_block_vip()

divider()

H2("Biến thể 11.3: Vé VVIP — Coaching Booking + Starter Kit", color=GREEN)
label("PHẦN COACHING")
body("📅 Đặt lịch coaching 1-1 với Lê Đăng Khương")
body("Đây là quyền lợi cao nhất của vé VVIP. 30 phút Zoom riêng với Thầy Khương "
     "— phân tích case của bạn, lộ trình áp dụng AI riêng, gỡ vướng mắc sâu.")
body("")
body("📌 Link Calendly đặt lịch: sẽ được gửi qua email riêng trong 24h. "
     "Lịch có giới hạn — book sớm để chọn được khung giờ ưng ý.")
body("")
body("→ Nút (vàng): 'Tải AI Agent Starter Kit ngay'  ·  Link: /courses/ai-make-more-money-vvip")
body("Bao gồm Template Database, Prompt Library & hướng dẫn dựng AI Agent đầu tiên trong 1 ngày.")
body("")

body("")
label("KẾT THƯ (chung)")
body("Cảm ơn bạn đã tin tưởng và đồng hành cùng KOHADA 💛")
body("— Lê Đăng Khương")

doc.add_page_break()

# ════════════════════════════════════════════════════════════════
# APPENDIX
# ════════════════════════════════════════════════════════════════
H1("📋 Hướng dẫn anh review & sửa", color=BLACK)
body("")
H3("Cách anh sửa nội dung")
bullet("Mở file Word, sửa trực tiếp các đoạn text — KHÔNG cần sửa cấu trúc/style")
bullet("Có thể đổi: subject, lời mở thư, lời kết, CTA copy, callout, bullet points, nội dung từng buổi")
bullet("KHÔNG đổi: link Zalo, route URL (/courses/*), placeholder {tên_khách}")
bullet("Gửi lại file Word cho em qua chat — em sẽ tự convert thành code và push deploy")
body("")
H3("Về placeholder {tên_khách}")
body("Trong tài liệu này em viết {tên_khách} chỉ để đánh dấu nơi tên thật xuất hiện. "
     "Email thật KHÔNG có chữ '{tên_khách}' — code tự inject tên thật từ database. "
     "Anh không cần thay thế hoặc xoá placeholder này.")
body("")
H3("Cách chạy event cho cohort mới")
body("Khi anh mở cohort tiếp theo (ví dụ tháng sau), chỉ cần:")
bullet("Update ngày trong file src/app/api/cron/aimakemoremoney/route.ts (SCHEDULE const)")
bullet("Nội dung email tự động dùng lại — không phải sửa gì")
bullet("Bảng aimm_attendees có thể reset hoặc giữ lại lịch sử (tuỳ ý)")
body("")
H3("Test trước khi event diễn ra")
body("Anh dùng email khác đăng ký thử ngay sau khi review xong, sẽ nhận welcome email "
     "trong 1 phút. Nếu nội dung đúng ý là OK.")

# Save
import os
os.makedirs(os.path.dirname(OUTPUT), exist_ok=True)
doc.save(OUTPUT)
print(f"✅ Saved: {OUTPUT}")
print(f"   File size: {os.path.getsize(OUTPUT) // 1024} KB")
